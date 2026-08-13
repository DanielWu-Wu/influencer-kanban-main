from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[2]
WORK = ROOT / ".codex-tmp" / "share-doc"
OUTPUT = ROOT / "docs" / "红人推广工作流看板_团队分享会讲稿.docx"

SKILL_SCRIPTS = Path(
    r"C:\Users\Admin\.codex\plugins\cache\openai-primary-runtime\documents\26.812.11052\skills\documents\scripts"
)
sys.path.insert(0, str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry  # noqa: E402


# launch_messaging_guide / compact_reference_guide tokens
PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
MARGIN_DXA = 1440
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 100, "bottom": 100, "start": 120, "end": 120}

FONT_CN = "Microsoft YaHei"
FONT_EN = "Aptos"
NAVY = "16324F"
BLUE = "2E74B5"
TEAL = "128C8C"
INK = "233143"
MUTED = "667085"
LIGHT_BLUE = "EAF2F8"
LIGHT_TEAL = "E8F6F4"
LIGHT_AMBER = "FFF5DA"
LIGHT_GRAY = "F3F5F7"
WHITE = "FFFFFF"
BORDER = "D7E0E8"
RED = "A33A3A"


def set_run_font(run, size=None, color=INK, bold=None, italic=None, name=FONT_CN):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color=BORDER, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, side="left", color=BLUE, size="18", space="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    edge = OxmlElement(f"w:{side}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), size)
    edge.set(qn("w:space"), space)
    edge.set(qn("w:color"), color)
    p_bdr.append(edge)


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def add_numbering_definitions(doc):
    numbering = doc.part.numbering_part.element
    existing_abs = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(existing_abs or [0]) + 1
    num_ids = []

    for kind in ("bullet", "decimal"):
        abs_num = OxmlElement("w:abstractNum")
        abs_num.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abs_num.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), kind)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
        lvl.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl.append(suff)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "270")
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), FONT_CN)
        fonts.set(qn("w:hAnsi"), FONT_CN)
        fonts.set(qn("w:eastAsia"), FONT_CN)
        r_pr.append(fonts)
        lvl.append(r_pr)
        abs_num.append(lvl)
        numbering.append(abs_num)

        num_id = max(existing_num + num_ids or [0]) + 1
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abs_ref = OxmlElement("w:abstractNumId")
        abs_ref.set(qn("w:val"), str(abstract_id))
        num.append(abs_ref)
        numbering.append(num)
        num_ids.append(num_id)
        abstract_id += 1

    return num_ids[0], num_ids[1]


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)


def add_bullet(doc, text, bullet_num_id, *, bold_prefix=None, color=INK):
    p = doc.add_paragraph()
    apply_num(p, bullet_num_id)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, 11, color=color, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, 11, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, 11, color=color)
    return p


def add_number(doc, text, decimal_num_id, *, bold_prefix=None):
    p = doc.add_paragraph()
    apply_num(p, decimal_num_id)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, 11, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, 11)
    else:
        r = p.add_run(text)
        set_run_font(r, 11)
    return p


def add_body(doc, text, *, bold_lead=None, color=INK, after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, 11, color=color, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2, 11, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, 11, color=color)
    return p


def add_callout(doc, label, text, *, fill=LIGHT_BLUE, accent=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.14)
    p.paragraph_format.right_indent = Inches(0.05)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.2
    set_paragraph_shading(p, fill)
    set_paragraph_border(p, color=accent)
    r1 = p.add_run(label + "  ")
    set_run_font(r1, 10.5, color=accent, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, 10.5, color=INK)
    return p


def add_h1(doc, text, kicker=None):
    if kicker:
        p0 = doc.add_paragraph()
        p0.paragraph_format.space_before = Pt(0)
        p0.paragraph_format.space_after = Pt(4)
        r0 = p0.add_run(kicker.upper())
        set_run_font(r0, 8.5, color=TEAL, bold=True, name=FONT_EN)
        set_keep_with_next(p0)
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.page_break_before = False
    p.add_run(text)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    p.add_run(text)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph(style="Heading 3")
    p.add_run(text)
    return p


def add_table(doc, headers, rows, widths, *, header_fill=NAVY, compact=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, header_fill)
        set_cell_borders(cell, color=header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        r = p.add_run(header)
        set_run_font(r, 9.5 if compact else 10, color=WHITE, bold=True)
    set_repeat_table_header(table.rows[0])
    for r_idx, row_values in enumerate(rows):
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
        for c_idx, value in enumerate(row_values):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_shading(cell, WHITE if r_idx % 2 == 0 else "F8FAFC")
            set_cell_borders(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            if c_idx == 0 and len(headers) <= 3:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_run_font(r, 9.2 if compact else 9.6, color=INK, bold=(c_idx == 0))
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=CONTENT_WIDTH_DXA,
        indent_dxa=TABLE_INDENT_DXA,
        cell_margins_dxa=CELL_MARGINS,
    )
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_after = Pt(2)
    return table


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_field(paragraph, field_code):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for el in (fld_begin, instr, fld_sep, text, fld_end):
        run._r.append(el)
    set_run_font(run, 8.5, color=MUTED)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_CN
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_CN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_CN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, NAVY, 10, 5),
    ):
        style = styles[name]
        style.font.name = FONT_CN
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_CN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_CN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_sections(doc):
    for section in doc.sections:
        section.page_width = Twips(PAGE_WIDTH_DXA)
        section.page_height = Twips(PAGE_HEIGHT_DXA)
        section.top_margin = Twips(MARGIN_DXA)
        section.bottom_margin = Twips(MARGIN_DXA)
        section.left_margin = Twips(MARGIN_DXA)
        section.right_margin = Twips(MARGIN_DXA)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        section.different_first_page_header_footer = True

        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("红人推广工作流看板  |  团队分享会讲稿")
        set_run_font(r, 8.2, color=MUTED)

        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("2026.08  ·  ")
        set_run_font(r, 8.5, color=MUTED)
        add_field(p, "PAGE")


def build_document():
    WORK.mkdir(parents=True, exist_ok=True)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    configure_styles(doc)
    configure_sections(doc)
    bullet_id, decimal_id = add_numbering_definitions(doc)

    # Cover: editorial_cover override for a live presentation guide.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(44)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("PROJECT STORY  ·  2026")
    set_run_font(r, 10, color=TEAL, bold=True, name=FONT_EN)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.keep_with_next = True
    r = p.add_run("红人推广工作流看板")
    set_run_font(r, 32, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(28)
    p.paragraph_format.keep_with_next = True
    r = p.add_run("从分散工具，到清晰、可控的海外红人运营工作台")
    set_run_font(r, 17, color=BLUE, bold=False)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    set_paragraph_border(p, side="bottom", color=TEAL, size="24", space="10")

    add_callout(
        doc,
        "一句话定位",
        "把 YouTube 红人发现、Gmail 沟通、飞书记录、合作履约、每日待办与 AI 辅助放进同一个桌面工作台；AI 帮忙，但所有关键外部动作仍由人确认。",
        fill=LIGHT_TEAL,
        accent=TEAL,
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("团队分享会讲稿 + 现场演示手册")
    set_run_font(r, 11, color=INK, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("建议时长：30–40 分钟  |  可压缩为 20–25 分钟")
    set_run_font(r, 10.5, color=MUTED)
    p = doc.add_paragraph()
    r = p.add_run("内容基于当前项目代码、README 与交接文档整理  |  2026 年 8 月")
    set_run_font(r, 9.5, color=MUTED)

    add_page_break(doc)

    # Page 2
    add_h1(doc, "先看这页：整场分享怎么讲", "RUN OF SHOW")
    add_callout(doc, "听众最后应记住", "这不是一个“又做了很多页面”的项目，而是对红人推广工作方式的一次重新组织：以业务流程为主线，以人工确认为安全边界，以 AI 为加速器。")
    add_table(
        doc,
        ["环节", "建议时间", "讲解重点", "是否演示"],
        [
            ("1. 为什么做", "5 分钟", "真实痛点与项目起点", "否"),
            ("2. 怎么思考", "8 分钟", "第一性原理、产品原则与关键取舍", "配合流程表"),
            ("3. 主要功能", "10 分钟", "按业务链路讲模块，不按菜单念功能", "轻演示"),
            ("4. 完整演示", "12–15 分钟", "从红人发现走到合作与每日跟进", "是"),
            ("5. 复盘与未来", "5 分钟", "当前边界、经验与下一步", "否"),
        ],
        [1600, 1350, 4300, 2110],
        compact=True,
    )
    add_h2(doc, "建议开场（可直接照着说）")
    add_body(doc, "大家好，今天我想分享的不是一段代码，而是我如何把一套真实的海外红人推广工作，重新设计成一个每天可以使用的工作台。")
    add_body(doc, "这项工作原来分散在 YouTube、Gmail、飞书和个人记忆里。问题不是缺少工具，而是工具之间没有一条连续的业务主线：红人在哪里发现、邮件跟到哪一步、样品是否发出、视频是否上线，常常需要来回切换才能拼出答案。")
    add_body(doc, "所以这个项目的核心，不是“再做一个 CRM”，而是让每一次关键动作更清楚、更安全、更容易追踪。接下来我会先讲设计过程，再用真实页面走一遍主要流程。")
    add_callout(doc, "讲者提示", "开场不要先介绍 Next.js、Supabase 或 API。先让听众理解业务问题，技术放到后面作为实现支撑。", fill=LIGHT_AMBER, accent="B27A00")

    add_page_break(doc)

    # Page 3
    add_h1(doc, "一、项目从什么问题开始", "PART 1 · WHY")
    add_h2(doc, "表面上工具很多，实际上流程仍靠人脑串联")
    add_table(
        doc,
        ["原来的工作场景", "带来的问题", "真正需要解决的事"],
        [
            ("YouTube 找红人", "频道资料、公开邮箱和最近内容要手工摘录", "让线索快速进入可跟进状态"),
            ("Gmail 沟通", "邮件、回复、Follow Up 和草稿状态容易断开", "看见沟通上下文与下一步动作"),
            ("飞书记记录", "字段多、表多，重复建档和漏写难避免", "保留结构化记录并减少错误写入"),
            ("合作履约", "发货、到货、拍摄、上线散在备注和记忆中", "让合作阶段、风险和时间点可见"),
            ("AI 辅助", "生成快，但可能误写、误发或覆盖人工判断", "让 AI 提建议，让人做决定"),
        ],
        [2600, 3100, 3660],
        compact=True,
    )
    add_h2(doc, "我把业务目标重新定义为三个词")
    add_h3(doc, "1. 清晰：当前状态一眼可见")
    add_body(doc, "运营人员不需要先回忆“我上次做了什么”，而应直接看到：这位红人在哪个阶段、最近一次沟通是什么、下一步应该做什么。")
    add_h3(doc, "2. 可控：每个外部动作有明确确认")
    add_body(doc, "发邮件、保存草稿、写飞书、更新合作状态都可能影响真实业务。系统可以准备内容、检查条件、展示预览，但最后一步必须由用户明确触发。")
    add_h3(doc, "3. 连续：流程不能停在“找到红人”")
    add_body(doc, "真正有价值的闭环，是从发现、建档、触达、合作、履约一直延伸到发布与复盘。项目的模块设计都围绕这条主线展开。")
    add_callout(doc, "一句总结", "问题不是“数据存在哪里”，而是“下一步是否清楚、动作是否安全、结果是否能回到流程里”。", fill=LIGHT_TEAL, accent=TEAL)

    add_page_break(doc)

    # Page 4
    add_h1(doc, "二、从第一性原理做出的设计选择", "PART 2 · THINKING")
    add_h2(doc, "原则 1：业务流程是主角，页面只是承载方式")
    add_body(doc, "我没有先问“要做几个页面”，而是先画出红人推广从开始到结束的状态变化。页面、按钮和字段都应该帮助状态向前推进，而不是让用户维护更多表格。")
    add_h2(doc, "原则 2：让专业系统各司其职")
    add_bullet(doc, "YouTube 负责公开频道与内容信息；看板不承诺获取隐藏邮箱。", bullet_id)
    add_bullet(doc, "Gmail 仍是邮件事实来源；系统负责上下文、起草、检查与草稿，不擅自发送。", bullet_id)
    add_bullet(doc, "飞书继续承载结构化业务数据；系统通过用户保存的字段映射读写，不硬编码字段名。", bullet_id)
    add_bullet(doc, "AI 负责判断、提取、翻译和生成；不能越过人工确认直接改变外部系统。", bullet_id)
    add_h2(doc, "原则 3：自动化应减少重复劳动，而不是减少人的知情权")
    add_body(doc, "例如 Follow Up 可以后台生成，但生成不等于保存；中文修改后必须重新同步外文；Gmail 草稿成功但飞书写回失败时，只重试飞书，不重复创建草稿。这样的设计看起来多一步，却能防止真实业务中的重复发送和错误覆盖。")
    add_h2(doc, "原则 4：多人协作的前提是账号隔离")
    add_body(doc, "产品、待办、Gmail、飞书授权、AI Key 和缓存都需要绑定当前账号。管理员只能管理账号状态，不能借此读取成员的邮件、业务数据或密钥。")
    add_callout(doc, "设计取舍", "这个项目优先选择“可解释、可预览、可恢复”，而不是追求完全自动化。对 B2B 运营工具来说，信任比炫技更重要。", fill=LIGHT_AMBER, accent="B27A00")

    add_page_break(doc)

    # Page 5
    add_h1(doc, "三、一条完整业务链路", "PART 2 · WORKFLOW")
    add_callout(doc, "整场分享的主轴", "后面介绍任何功能时，都回到一个问题：它在这条链路上解决哪一个断点？", fill=LIGHT_TEAL, accent=TEAL)
    add_table(
        doc,
        ["阶段", "主要动作", "系统提供的帮助"],
        [
            ("1. 发现", "在 YouTube 找到潜在频道", "识别公开资料、近期内容与邮箱线索"),
            ("2. 建档", "判断是否值得联系并进入资源库", "飞书双表查重、补全和人工确认"),
            ("3. 触达", "准备首封开发信与 Follow Up", "产品上下文、双语起草、Gmail 草稿"),
            ("4. 合作", "确认方式、价格、样品与时间", "邮件上下文、合作记录和关键条件"),
            ("5. 履约", "发货、到货、拍摄与发布", "阶段推导、风险提示、告知与日历节点"),
            ("6. 复盘", "沉淀内容表现与合作经验", "当前逐步补齐，未来连接数据闭环"),
        ],
        [1450, 3400, 4510],
        compact=True,
    )
    add_h2(doc, "系统之间的分工")
    add_table(
        doc,
        ["系统", "主要职责", "工作台如何衔接"],
        [
            ("YouTube", "公开频道与内容信息", "把频道信息转成可查重、可建档的红人线索"),
            ("Gmail", "真实邮件、线程与草稿", "汇总上下文、辅助起草，并保留人工发送"),
            ("飞书", "结构化业务记录", "依照用户字段映射读取与确认写回"),
            ("AI", "判断、提取、翻译与生成", "提供建议和内容，不越过用户执行外部动作"),
            ("工作台", "聚合上下文并推动流程", "让下一步、风险、确认点和结果可见"),
        ],
        [1450, 3500, 4410],
        compact=True,
    )
    add_callout(doc, "讲者提示", "讲到这里可停 20 秒：看板不是替代 YouTube、Gmail 或飞书，而是把它们连接成可推进的工作流。", fill=LIGHT_BLUE, accent=BLUE)

    add_page_break(doc)

    # Page 6
    add_h1(doc, "四、主要功能：红人开发台", "PART 3 · FEATURES")
    add_callout(doc, "它解决什么", "把“找到一个频道”变成“得到一条可联系、已查重、可进入开发信流程的线索”。")
    add_h2(doc, "四个连续步骤")
    add_number(doc, "红人录入：批量粘贴 YouTube 链接、@handle 或频道 ID，识别频道、头像、国家、语言、公开邮箱与最近长视频。", decimal_id, bold_prefix="红人录入：")
    add_number(doc, "邀约确认：选择目标产品、合作形式、合作想法、优先级和开发信语言，缺少关键条件时不进入下一步。", decimal_id, bold_prefix="邀约确认：")
    add_number(doc, "开发信：AI 根据频道与产品上下文生成外文正文、中文对照和备选标题；用户检查后只保存 Gmail 草稿。", decimal_id, bold_prefix="开发信：")
    add_number(doc, "开发信跟进：根据真实 Gmail 状态区分真人回复、自动回复、退信与待跟进；单条或批量生成后逐封审核、逐封保存。", decimal_id, bold_prefix="开发信跟进：")
    add_h2(doc, "我最看重的三个细节")
    add_bullet(doc, "先查重再建档：避免同一红人在资源库和开发记录里重复出现。", bullet_id)
    add_bullet(doc, "多邮箱不替用户做决定：系统合并来源并提示选择，人工修改后不再自动覆盖。", bullet_id)
    add_bullet(doc, "公开信息有边界：只能从公开简介提取邮箱，无法获取频道隐藏的商务邮箱。", bullet_id)
    add_h2(doc, "现场演示建议（3–4 分钟）")
    add_table(
        doc,
        ["操作", "你可以边操作边讲"],
        [
            ("打开“红人开发台”", "四个 Tab 不是四个孤立工具，而是同一条线索的四个阶段。"),
            ("粘贴一个测试频道", "系统先补全公开资料，再做飞书查重；没有邮箱时不会假装找到。"),
            ("进入邀约确认", "让 AI 有足够业务上下文，避免只凭频道名写一封泛化邮件。"),
            ("展示邮件预览", "强调中文对照、产品资料和“只保存草稿，不自动发送”。"),
        ],
        [2350, 7010],
        compact=True,
    )

    add_page_break(doc)

    # Page 7
    add_h1(doc, "五、主要功能：Gmail 与 AI 辅助回复", "PART 3 · FEATURES")
    add_callout(doc, "它解决什么", "把邮件内容、联系人上下文、红人资料、翻译与下一步动作放在同一个阅读和起草界面中。")
    add_h2(doc, "核心能力")
    add_bullet(doc, "查看收件箱、未读、星标、已发送与草稿，并使用 Gmail 全局搜索查找较早邮件。", bullet_id)
    add_bullet(doc, "可选择线程中的任意真实邮件作为回复锚点，查看 From / To / Cc，并在人选不明确时要求确认收件人。", bullet_id)
    add_bullet(doc, "AI 回复会结合最近往来、红人资料与模板规则生成外文和中文对照；画像优化版不会自动覆盖当前草稿。", bullet_id)
    add_bullet(doc, "用户修改中文或目标语言后，需要重新翻译；外文手工润色可保留，避免 AI 抹掉人工表达。", bullet_id)
    add_bullet(doc, "已发送邮件默认拦截第三方远程图片和样式资源，降低看板查看触发 Mailsuite / Mailtrack 打开记录的风险。", bullet_id)
    add_h2(doc, "为什么“生成”和“执行”必须分开")
    add_table(
        doc,
        ["AI 可以做", "AI 不应擅自做"],
        [
            ("总结线程、识别语言、起草和翻译", "替用户确认预算、日期和合作条件"),
            ("提供回复版本与画像优化建议", "自动覆盖用户正在编辑的草稿"),
            ("检查收件人与模板条件", "自动发送邮件或批量群发"),
            ("准备飞书写回预览", "绕过字段映射直接修改外部数据"),
        ],
        [4680, 4680],
        compact=True,
    )
    add_h2(doc, "现场演示建议（4–5 分钟）")
    add_number(doc, "打开一条已有测试线程，先展示参与者、红人资料和历史上下文。", decimal_id)
    add_number(doc, "点击 AI 辅助回复，展示流式生成与中文对照；说明画像分析不会自动改稿。", decimal_id)
    add_number(doc, "修改一小段中文，让听众看到必须重新同步外文的安全门。", decimal_id)
    add_number(doc, "停在草稿预览或保存确认框；除非使用专门测试账号，不要现场发送。", decimal_id)

    add_page_break(doc)

    # Page 8
    add_h1(doc, "六、主要功能：合作项目、每日待办与工作日历", "PART 3 · FEATURES")
    add_h2(doc, "合作项目：把“联系上了”继续推进到“真正交付了”")
    add_body(doc, "合作项目读取飞书详细合作记录，提供列表、看板、日历、搜索、阶段筛选和风险提示。阶段由确认、发货、到货、拍摄、预计上线和实际上线等业务字段推导，而不是只靠一个手工状态。")
    add_bullet(doc, "在项目详情中查看关键日期、物流和折扣告知状态。", bullet_id)
    add_bullet(doc, "物流或折扣邮件先生成、编辑和翻译，再保存到最近的正常 Gmail 线程。", bullet_id)
    add_bullet(doc, "草稿保存成功后才同步飞书“已告知”；若飞书失败，只重试写回，不重复建草稿。", bullet_id)

    add_h2(doc, "每日待办：把“今天最该做什么”放到首页")
    add_body(doc, "待办既包括手动任务，也会读取近 72 小时符合条件的 Gmail 外部来信，排除推广、自动回复、退信和追踪通知，并匹配飞书中的红人。用户回复后任务可自动完成，新来信出现时会重新进入待完成。")

    add_h2(doc, "工作日历：把任务与合作节点放进同一个时间视图")
    add_body(doc, "工作日历汇总手动日程、未完成待办和合作项目六类节点。合作节点只读，点击后跳转到项目详情；逾期且未发布的预计上线节点会突出显示。")
    add_callout(doc, "设计重点", "这三个模块共同解决“合作中段和后段容易失控”的问题：不是记录更多，而是更早看见风险和下一步。", fill=LIGHT_TEAL, accent=TEAL)
    add_h2(doc, "现场演示建议（4–5 分钟）")
    add_table(
        doc,
        ["页面", "演示动作", "重点话术"],
        [
            ("合作项目", "切换列表 / 看板 / 日历，打开一条项目详情", "阶段来自真实业务字段，风险与时间点可见。"),
            ("每日待办", "展示待完成、今日已完成、历史已完成", "系统帮助我确定今天先做什么。"),
            ("工作日历", "点击一个合作节点并跳转项目", "任务、日程和履约节点在时间上被连接起来。"),
        ],
        [1500, 3150, 4710],
        compact=True,
    )

    add_page_break(doc)

    # Page 9
    add_h1(doc, "七、支撑整个工作台的基础能力", "PART 3 · FOUNDATION")
    add_table(
        doc,
        ["模块", "作用", "设计边界"],
        [
            ("产品资料", "为开发信和合作邮件提供产品名称、型号、链接、卖点与主图", "沿用现有资源字段，不为主图额外引入存储系统"),
            ("飞书设置", "配置三张业务表、字段映射、只读检查与样例预览", "不硬编码字段名，不自动改表结构"),
            ("AI 设置与提示词", "允许用户配置 OpenAI-compatible 模型与不同起草规则", "项目不把某个模型供应商描述为内置服务"),
            ("账号管理", "管理员创建、停用、恢复账号并设置临时密码", "管理员看不到成员密码、邮件、业务数据与密钥"),
            ("缓存与反馈", "减少重复读取，让线程和页面切换更稳定", "缓存按账号隔离，失败要给出可理解的重试提示"),
        ],
        [1500, 4000, 3860],
        compact=True,
    )
    add_h2(doc, "给非技术听众的技术说明")
    add_body(doc, "前端采用 Next.js、React 与 TypeScript 构建桌面工作台；Supabase 负责登录与按账号隔离的数据；后端接口连接 Gmail、飞书、YouTube 和用户自行配置的 AI 模型服务。")
    add_body(doc, "技术选型不是分享重点，但有一个原则值得强调：敏感密钥只留在服务端，外部写入要经过用户确认，账号之间的数据和授权不能串用。")
    add_callout(doc, "可直接说", "我不希望技术成为新的风险源。这个项目不是为了把所有系统搬到一起，而是让它们在明确边界下协作。", fill=LIGHT_BLUE, accent=BLUE)
    add_h2(doc, "项目设计语言")
    add_bullet(doc, "桌面运营工作台：高信息密度，但层级清楚、可快速扫描。", bullet_id)
    add_bullet(doc, "轻量 Glassmorphism：用透明度、层次和柔和阴影建立空间感，不做花哨落地页。", bullet_id)
    add_bullet(doc, "反馈优先：加载、成功、失败、待确认和重试都应有明确状态。", bullet_id)
    add_bullet(doc, "保守默认：不确定时不静默写入，不用“智能”替代真实业务判断。", bullet_id)

    add_page_break(doc)

    # Page 10
    add_h1(doc, "八、现场演示脚本：用一条线索串起全项目", "PART 4 · LIVE DEMO")
    add_callout(doc, "演示目标", "让听众看到信息如何从“一个频道”流向“一个可跟进的合作项目”，而不是逐个菜单点一遍。", fill=LIGHT_TEAL, accent=TEAL)
    add_table(
        doc,
        ["时间", "操作路径", "讲解要点"],
        [
            ("0:00–1:00", "首页 / 每日待办", "先展示今天的工作入口：待办数、Gmail 来信与任务分组。"),
            ("1:00–4:00", "红人开发台 → 红人录入", "粘贴测试频道，展示公开资料识别、邮箱边界和飞书查重。"),
            ("4:00–7:00", "邀约确认 → 开发信", "补充产品和合作想法，生成双语邮件，强调只保存草稿。"),
            ("7:00–10:00", "Gmail 邮件", "打开测试线程，展示上下文、AI 回复、中文确认与回复锚点。"),
            ("10:00–13:00", "合作项目", "打开一条已合作项目，展示阶段、关键日期、告知邮件和风险。"),
            ("13:00–15:00", "工作日历", "从合作节点跳回项目，收束为“今天该做什么、项目走到哪”。"),
        ],
        [1550, 3050, 4760],
        compact=True,
    )
    add_h2(doc, "演示前 10 分钟检查")
    for item in [
        "准备一个不会影响真实业务的 YouTube 测试频道和一条已存在的测试合作项目。",
        "确认登录账号、Gmail、飞书、YouTube 与 AI 配置可用；提前打开页面并刷新一次。",
        "关闭无关通知，浏览器缩放恢复为 100%，侧边栏保持展开。",
        "如果要展示草稿保存，使用专门测试邮箱；默认不要现场发送邮件或修改真实飞书记录。",
        "准备两张备用截图或直接停留在本 Word 文档的功能页，以防网络或 OAuth 临时失效。",
    ]:
        add_bullet(doc, item, bullet_id)
    add_h2(doc, "演示时的节奏")
    add_body(doc, "每切换一个页面，只回答三个问题：现在处于哪个业务阶段？这个页面替用户省掉了什么？哪一步仍需要人确认？这样能避免演示变成按钮巡游。")

    add_page_break(doc)

    # Page 11
    add_h1(doc, "九、哪些内容可以放心讲，哪些要谨慎表述", "PART 4 · STATUS")
    add_table(
        doc,
        ["表达级别", "可讲内容", "推荐说法"],
        [
            ("已实现", "红人开发台、Gmail 辅助回复、合作项目、待办/日历、字段映射、账号隔离", "“当前版本已经提供……”"),
            ("需要现场条件", "真实 Gmail / 飞书读取、草稿保存、OAuth、多人账号", "“在配置完成并授权后可以……”"),
            ("仍需真实验收", "特定追踪场景、真实写回幂等、多成员端到端隔离", "“代码路径已具备，仍需用真实账号完成最终验收。”"),
            ("未来规划", "AI 识别写回信息后提醒、合作状态变更建议、发布后数据闭环", "“下一阶段计划……”"),
        ],
        [1700, 3900, 3760],
        compact=True,
    )
    add_h2(doc, "不要在分享会上这样说")
    add_bullet(doc, "不要说“AI 会自动管理整个合作”；应说“AI 会准备建议和内容，关键动作由人确认”。", bullet_id)
    add_bullet(doc, "不要说“YouTube API 能拿到红人邮箱”；应说“系统只能从公开简介提取可见邮箱”。", bullet_id)
    add_bullet(doc, "不要说“系统会自动发开发信”；应说“系统创建并检查草稿，用户在 Gmail 手动发送”。", bullet_id)
    add_bullet(doc, "不要把尚未部署或尚未用真实账号验证的能力描述为生产环境已验收。", bullet_id)
    add_h2(doc, "出现现场故障时怎么接")
    add_table(
        doc,
        ["情况", "处理方式"],
        [
            ("Gmail / 飞书授权过期", "不要现场反复登录；切回 Word，按对应功能页讲设计与安全边界。"),
            ("AI 响应较慢", "先讲输入上下文和人工确认规则；等待期间演示已有草稿或其他页面。"),
            ("频道没有公开邮箱", "这是正常结果，正好说明公开数据边界；换预备频道继续。"),
            ("写回失败", "展示错误和重试设计，不要连续点击；强调不会重复创建草稿。"),
        ],
        [2800, 6560],
        compact=True,
    )

    add_page_break(doc)

    # Page 12
    add_h1(doc, "十、这次项目最重要的设计收获", "PART 5 · REFLECTION")
    add_h2(doc, "1. 从“做功能”转向“设计工作流”")
    add_body(doc, "当我围绕业务状态设计时，很多页面关系自然清楚了：红人开发台负责前半段，合作项目负责履约，每日待办和工作日历负责把跨模块的下一步重新聚合。")
    add_h2(doc, "2. AI 的价值来自上下文与边界")
    add_body(doc, "单纯生成一封邮件并不难；真正困难的是让 AI 理解频道、产品、历史邮件与合作条件，同时不覆盖人工信息、不选错收件人、不越权发送。")
    add_h2(doc, "3. 可靠性体现在失败时")
    add_body(doc, "好的工作台不仅要在成功时快，还要在 Gmail 成功、飞书失败，或网络、授权、缓存异常时告诉用户发生了什么、能否重试、会不会重复执行。")
    add_h2(doc, "4. 产品成熟度不是功能数量")
    add_body(doc, "对运营工具来说，成熟度来自数据准确、状态可见、权限清楚、操作可恢复。很多“保守”的设计，恰恰是在真实业务里建立信任的关键。")
    add_callout(doc, "我的结论", "这个项目最有价值的部分，不是把 AI 塞进 CRM，而是重新定义了人、AI 与外部系统之间如何分工。", fill=LIGHT_TEAL, accent=TEAL)

    add_h2(doc, "下一阶段：从运营工作台走向数据闭环")
    add_table(
        doc,
        ["方向", "目标", "仍保持的边界"],
        [
            ("合作履约深化", "多产品筛选、内容进度、发布资料与风险提醒", "不靠模糊字符串匹配，不自动改状态"),
            ("AI 写回建议", "识别邮件中的报价、地址、档期和合作状态变化", "先生成变更预览，再由用户确认写回"),
            ("发布后数据", "连接视频表现、联盟与销售转化，支持复盘", "公开数据与估算数据分开标注"),
            ("团队协作", "进一步完善成员配置与任务协同", "账号、授权、密钥和业务数据继续隔离"),
        ],
        [1900, 4100, 3360],
        compact=True,
    )

    add_page_break(doc)

    # Page 13
    add_h1(doc, "十一、收尾与问答", "PART 5 · CLOSE")
    add_h2(doc, "建议收尾（可直接照着说）")
    add_body(doc, "回到最开始的问题：我们缺的并不是更多工具，而是一条清楚、安全、能持续推进的工作流。这个项目把红人发现、邮件沟通、合作履约和每日执行连接起来，也让我更明确地看到 AI 在业务系统中的正确位置——它应该提高判断和表达的效率，但不能替人承担未经确认的业务决策。")
    add_body(doc, "现在这个工作台已经覆盖了从红人开发到合作跟进的主要环节。下一步，我希望继续补足发布后的数据反馈，让每次合作的经验可以沉淀到下一次选择和沟通中。谢谢大家，接下来欢迎提问。")

    add_h2(doc, "听众可能会问")
    qa_rows = [
        ("为什么不直接在飞书里完成？", "飞书适合结构化记录，但邮件上下文、AI 起草、实时任务和跨系统确认需要更连贯的交互；工作台保留飞书作为数据来源，而不是替代它。"),
        ("为什么不让 AI 自动发送？", "邮件和飞书写回会产生真实业务后果。当前阶段优先保证可预览、可确认、可恢复，自动化只做到低风险准备环节。"),
        ("这个项目最大的难点是什么？", "不是页面数量，而是跨系统状态一致、收件人安全、字段映射、账号隔离以及失败后的幂等重试。"),
        ("后续最值得做什么？", "先把合作履约和发布后数据闭环补齐，再让 AI 基于真实事件提出可确认的运营建议。"),
        ("团队成员能看到彼此的数据吗？", "账号业务数据、Gmail、飞书授权与密钥按账号隔离；管理员只管理账号元数据，不能读取成员内容。"),
    ]
    add_table(doc, ["问题", "回答要点"], qa_rows, [2900, 6460], compact=True)
    add_callout(doc, "最后提醒", "正式分享前至少完整走一遍演示路径；如果任何真实写入步骤没有测试账号，就停在预览或确认框。", fill=LIGHT_AMBER, accent="B27A00")

    add_h2(doc, "材料依据")
    add_body(doc, "本讲稿依据当前项目 README、docs/HANDOFF.md、主导航与红人开发、Gmail、合作项目、每日待办、工作日历、设置和账号管理相关代码整理。功能状态以分享当天真实运行环境为准。", color=MUTED, after=3)

    # Document properties
    doc.core_properties.title = "红人推广工作流看板｜团队分享会讲稿"
    doc.core_properties.subject = "项目设计思考、主要功能与现场演示脚本"
    doc.core_properties.author = "项目团队"
    doc.core_properties.keywords = "红人推广, CRM, Gmail, 飞书, YouTube, AI, 分享会"

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
